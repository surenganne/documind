# Feature: documind-platform, Property 6: Document Text Extraction Produces Non-Empty Content
"""
Property 6: For any valid PDF, DOCX, TXT, or Markdown file, text extraction
produces a non-empty string. For DOCX files, heading hierarchy is preserved.
"""
import io
import os
import tempfile
import textwrap

import pytest
from hypothesis import given, settings, strategies as st

from app.services.document.extractor import extract_text


# ── Helpers to build synthetic files ─────────────────────────────────────────

def _make_txt(content: str) -> str:
    """Write content to a temp .txt file and return its path."""
    f = tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8")
    f.write(content)
    f.close()
    return f.name


def _make_md(content: str) -> str:
    f = tempfile.NamedTemporaryFile(suffix=".md", delete=False, mode="w", encoding="utf-8")
    f.write(content)
    f.close()
    return f.name


def _make_docx(headings: list[tuple[int, str]], body_lines: list[str]) -> str:
    """
    Build a minimal DOCX with the given headings and body paragraphs.
    headings: list of (level, text) e.g. [(1, "Introduction"), (2, "Background")]
    """
    from docx import Document as DocxDocument

    doc = DocxDocument()
    for level, text in headings:
        doc.add_heading(text, level=level)
    for line in body_lines:
        doc.add_paragraph(line)

    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    f.close()
    doc.save(f.name)
    return f.name


def _make_pdf(text_content: str) -> str:
    """Build a minimal single-page PDF using reportlab."""
    try:
        from reportlab.pdfgen import canvas as rl_canvas
    except ImportError:
        pytest.skip("reportlab not installed — skipping PDF synthesis")

    f = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    f.close()
    c = rl_canvas.Canvas(f.name)
    # Write text line by line (max ~80 chars per line to stay on page)
    y = 750
    for line in textwrap.wrap(text_content, width=80) or [text_content]:
        c.drawString(50, y, line)
        y -= 15
        if y < 50:
            c.showPage()
            y = 750
    c.save()
    return f.name


# ── Strategies ────────────────────────────────────────────────────────────────

# Non-empty printable text (avoid pure-whitespace strings)
nonempty_text = st.text(
    alphabet=st.characters(whitelist_categories=("Lu", "Ll", "Nd", "Zs", "Po")),
    min_size=5,
    max_size=200,
).filter(lambda s: s.strip())

heading_level = st.integers(min_value=1, max_value=4)
heading_text = st.text(
    alphabet=st.characters(whitelist_categories=("Lu", "Ll", "Nd", "Zs")),
    min_size=3,
    max_size=40,
).filter(lambda s: s.strip())


# ── TXT tests ─────────────────────────────────────────────────────────────────

@given(nonempty_text)
@settings(max_examples=50)
def test_txt_extraction_nonempty(content: str):
    """Property 6 (TXT): extraction of any non-empty TXT file yields non-empty string."""
    path = _make_txt(content)
    try:
        result = extract_text(path, "txt")
        assert result.strip(), f"Expected non-empty extraction, got: {result!r}"
    finally:
        os.unlink(path)


@given(nonempty_text)
@settings(max_examples=50)
def test_md_extraction_nonempty(content: str):
    """Property 6 (MD): extraction of any non-empty Markdown file yields non-empty string."""
    path = _make_md(content)
    try:
        result = extract_text(path, "md")
        assert result.strip(), f"Expected non-empty extraction, got: {result!r}"
    finally:
        os.unlink(path)


def test_txt_extraction_preserves_content():
    """Extracted TXT content contains the original text."""
    path = _make_txt("Hello DocuMind world")
    try:
        result = extract_text(path, "txt")
        assert "Hello DocuMind world" in result
    finally:
        os.unlink(path)


# ── DOCX tests ────────────────────────────────────────────────────────────────

@given(
    st.lists(st.tuples(heading_level, heading_text), min_size=1, max_size=5),
    st.lists(nonempty_text, min_size=1, max_size=3),
)
@settings(max_examples=30)
def test_docx_extraction_nonempty(headings, body_lines):
    """Property 6 (DOCX): extraction of any DOCX with headings yields non-empty string."""
    path = _make_docx(headings, body_lines)
    try:
        result = extract_text(path, "docx")
        assert result.strip(), f"Expected non-empty extraction, got: {result!r}"
    finally:
        os.unlink(path)


@given(
    st.lists(st.tuples(heading_level, heading_text), min_size=1, max_size=4),
)
@settings(max_examples=30)
def test_docx_preserves_heading_hierarchy(headings):
    """
    Property 6 (DOCX heading hierarchy): for each heading in the DOCX,
    the extracted text must contain a markdown-style prefix (# / ## / ###)
    matching the heading level.
    """
    path = _make_docx(headings, ["Some body text."])
    try:
        result = extract_text(path, "docx")
        for level, text in headings:
            expected_prefix = "#" * level
            # Normalize whitespace for comparison (Unicode spaces may differ)
            normalized_text = " ".join(text.split())
            # The heading text should appear with the correct # prefix
            assert any(
                line.startswith(expected_prefix)
                and " ".join(line.split()).find(normalized_text) >= 0
                for line in result.splitlines()
            ), (
                f"Heading level {level} '{text}' not found with prefix '{expected_prefix}' in:\n{result}"
            )
    finally:
        os.unlink(path)


def test_docx_heading_levels_distinct():
    """H1 and H2 headings must produce different prefix depths."""
    path = _make_docx([(1, "TopLevel"), (2, "SubLevel")], [])
    try:
        result = extract_text(path, "docx")
        lines = result.splitlines()
        h1_lines = [l for l in lines if l.startswith("# ") and not l.startswith("## ")]
        h2_lines = [l for l in lines if l.startswith("## ")]
        assert any("TopLevel" in l for l in h1_lines), "H1 heading not found"
        assert any("SubLevel" in l for l in h2_lines), "H2 heading not found"
    finally:
        os.unlink(path)


# ── PDF tests ─────────────────────────────────────────────────────────────────

@given(nonempty_text)
@settings(max_examples=20)
def test_pdf_extraction_nonempty(content: str):
    """Property 6 (PDF): extraction of any non-empty PDF yields non-empty string."""
    path = _make_pdf(content)
    try:
        result = extract_text(path, "pdf")
        assert result.strip(), f"Expected non-empty extraction, got: {result!r}"
    finally:
        os.unlink(path)


# ── Error handling ────────────────────────────────────────────────────────────

def test_unsupported_file_type_raises():
    """Unsupported file types must raise ValueError."""
    path = _make_txt("test")
    try:
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(path, "xlsx")
    finally:
        os.unlink(path)


def test_missing_file_raises():
    """Non-existent file path must raise FileNotFoundError."""
    with pytest.raises(FileNotFoundError):
        extract_text("/tmp/does_not_exist_12345.txt", "txt")
